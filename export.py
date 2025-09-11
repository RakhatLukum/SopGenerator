from __future__ import annotations
import os
import re
from datetime import datetime
from typing import Dict, List, Tuple
from io import BytesIO

from docx import Document
from docx.shared import Pt, Mm, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT


def _set_document_defaults(doc: Document, font_name: str = "Times New Roman", font_size_pt: int = 12) -> None:
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)
    # Ensure East Asia font mapping
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn("w:eastAsia"), font_name)

    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    # 1 inch margins
    margin = Inches(1)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin


def _add_page_numbers(section) -> None:
    # Add "Page X of Y" to footer using field codes
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_field_run(p, instr_text: str):
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), instr_text)
        r = OxmlElement("w:r")
        fld.append(r)
        p._p.append(fld)

    run = paragraph.add_run("Page ")
    _add_field_run(paragraph, "PAGE")
    run = paragraph.add_run(" of ")
    _add_field_run(paragraph, "NUMPAGES")


def _add_title_page(doc: Document, meta: Dict) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run((meta.get("title") or "Стандартная операционная процедура").strip())
    run.bold = True
    run.font.size = Pt(20)

    doc.add_paragraph("")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Стандартная операционная процедура (СОП)")
    run2.font.size = Pt(14)

    doc.add_paragraph("")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f"Дата: {datetime.now().strftime('%Y-%m-%d')}")
    run3.font.size = Pt(10)

    doc.add_page_break()


def _add_approval_sheet(doc: Document, meta: Dict) -> None:
    doc.add_heading("Approval Sheet", level=1)
    table = doc.add_table(rows=4, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Role"
    hdr_cells[1].text = "Name"
    hdr_cells[2].text = "Signature/Date"

    rows = [
        ("Prepared by", meta.get("prepared_by", ""), "________________ / __________"),
        ("Reviewed by", meta.get("reviewed_by", ""), "________________ / __________"),
        ("Approved by", meta.get("approved_by", ""), "________________ / __________"),
    ]
    for i, (role, name, sign) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = role
        cells[1].text = name
        cells[2].text = sign

    doc.add_paragraph("")


def _add_change_log(doc: Document, meta: Dict) -> None:
    doc.add_heading("Журнал изменений", level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Версия"
    hdr[1].text = "Дата"
    hdr[2].text = "Описание"
    hdr[3].text = "Автор"

    changes: List[Dict] = meta.get("changes", []) or []
    if not changes:
        changes = [
            {"version": meta.get("version", 1), "date": datetime.now().strftime("%Y-%m-%d"), "description": "Первоначальная версия", "author": meta.get("author", "Автор")}
        ]
    for ch in changes:
        row = table.add_row().cells
        row[0].text = str(ch.get("version", ""))
        row[1].text = str(ch.get("date", ""))
        row[2].text = str(ch.get("description", ""))
        row[3].text = str(ch.get("author", ""))

    doc.add_paragraph("")


def _add_acknowledgement(doc: Document, meta: Dict) -> None:
    doc.add_heading("Благодарности", level=1)
    text = meta.get("acknowledgement") or "Мы благодарим всех участников за вклад в разработку данной СОП."
    doc.add_paragraph(text)


def _iter_markdown_lines(md: str):
    for line in (md or "").splitlines():
        yield line.rstrip("\n")


def _is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|?\s*:?[-]+:?\s*(\|\s*:?[-]+:?\s*)+\|?\s*$", line))


def _parse_markdown_table(lines: List[str], idx: int) -> Tuple[List[List[str]], int]:
    # Very basic GitHub-style table parser
    header = [c.strip() for c in lines[idx].strip().strip('|').split('|')]
    idx += 1
    # Skip separator
    if idx < len(lines) and _is_table_separator(lines[idx]):
        idx += 1
    rows: List[List[str]] = [header]
    while idx < len(lines):
        line = lines[idx]
        if not line.strip().startswith('|'):
            break
        rows.append([c.strip() for c in line.strip().strip('|').split('|')])
        idx += 1
    return rows, idx


def _add_table(doc: Document, rows: List[List[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(rows[0]):
        hdr_cells[i].text = h
    for r in rows[1:]:
        row_cells = table.add_row().cells
        for i, c in enumerate(r):
            row_cells[i].text = c


def _convert_markdown_to_doc(doc: Document, md: str) -> None:
    # Simple Markdown to DOCX conversion with numbering and basic lists
    lines = list(_iter_markdown_lines(md))
    sec_nums: List[int] = []  # hierarchical numbering
    table_counter = 0
    figure_counter = 0

    i = 0
    while i < len(lines):
        line = lines[i]

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # update counters
            while len(sec_nums) < level:
                sec_nums.append(0)
            while len(sec_nums) > level:
                sec_nums.pop()
            sec_nums[-1] += 1
            # reset deeper levels if any remain
            for j in range(level, len(sec_nums)):
                if j > level - 1:
                    sec_nums[j] = 0
            number = ".".join(str(n) for n in sec_nums if n > 0)
            # Strip existing numeric prefixes to avoid double numbering
            text = re.sub(r"^\d+(?:\.\d+)*\.?\s+", "", text)
            p = doc.add_paragraph()
            p.style = f"Heading {min(level, 3)}"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{number} {text}")
            run.bold = True
            i += 1
            continue

        # Tables
        if line.strip().startswith('|') and i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
            rows, i = _parse_markdown_table(lines, i)
            _add_table(doc, rows)
            table_counter += 1
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.add_run(f"Таблица {table_counter}: {rows[0][0] if rows and rows[0] else ''}").italic = True
            continue

        # Figure syntax ![caption](path)
        fm = re.match(r"^!\[(.*?)\]\((.*?)\)$", line)
        if fm:
            figure_counter += 1
            path = fm.group(2)
            caption = fm.group(1)
            try:
                if os.path.exists(path):
                    doc.add_picture(path, width=Inches(5.5))
            except Exception:
                pass
            pcap = doc.add_paragraph()
            pcap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pcap.add_run(f"Рисунок {figure_counter}: {caption}").italic = True
            i += 1
            continue

        # Lists
        if re.match(r"^\s*[-*]\s+", line):
            p = doc.add_paragraph(line.strip()[1:].strip(), style="List Bullet")
            i += 1
            continue
        if re.match(r"^\s*\d+[\.)]\s+", line):
            p = doc.add_paragraph(re.sub(r"^\s*\d+[\.)]\s+", "", line), style="List Number")
            i += 1
            continue

        # Footnotes (collect as plain text lines)
        if re.match(r"^\[\^\d+\]:", line):
            doc.add_paragraph(line)
            i += 1
            continue

        # Regular paragraph
        if line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
        i += 1


def export_docx(markdown_content: str, metadata: Dict, output_path: str) -> str:
    """Create a DOCX file with SOP formatting and return saved path."""
    doc = Document()
    _set_document_defaults(doc)

    _add_title_page(doc, metadata)
    _add_approval_sheet(doc, metadata)
    _add_change_log(doc, metadata)
    _add_acknowledgement(doc, metadata)

    doc.add_page_break()

    # Main content
    _convert_markdown_to_doc(doc, markdown_content)

    # Footer page numbers
    for section in doc.sections:
        _add_page_numbers(section)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def export_docx_bytes(markdown_content: str, metadata: Dict) -> bytes:
    """Create a DOCX in memory and return bytes for download."""
    doc = Document()
    _set_document_defaults(doc)

    _add_title_page(doc, metadata)
    _add_approval_sheet(doc, metadata)
    _add_change_log(doc, metadata)
    _add_acknowledgement(doc, metadata)

    doc.add_page_break()
    _convert_markdown_to_doc(doc, markdown_content)

    for section in doc.sections:
        _add_page_numbers(section)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue() 